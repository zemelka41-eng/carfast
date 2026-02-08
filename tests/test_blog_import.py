import re

import pytest
from pathlib import Path

from django.core.management import call_command

from blog.management.commands.import_blog_content import get_source_files, parse_docx_file
from blog.models import BlogPost


pytestmark = pytest.mark.django_db


def _import_posts(include_all: bool = True):
    if include_all:
        call_command("import_blog_content", "--all")
    else:
        call_command("import_blog_content")
    files, _ = get_source_files(include_all=include_all)
    data = [parse_docx_file(path) for path in files]
    titles = [item["title"] for item in data]
    return BlogPost.objects.filter(title__in=titles), titles, data


def test_blog_detail_returns_200_for_imported_posts(client):
    posts, _, _ = _import_posts(include_all=True)
    assert posts.count() >= 2
    for post in posts:
        response = client.get(post.get_absolute_url())
        assert response.status_code == 200


def test_blog_post_has_canonical_and_single_h1(client):
    posts, _, _ = _import_posts(include_all=True)
    for post in posts:
        response = client.get(post.get_absolute_url())
        content = response.content.decode("utf-8")
        assert f'rel="canonical" href="https://carfst.ru{post.get_absolute_url()}"' in content
        assert len(re.findall(r"<h1", content)) == 1


def test_blog_post_indexable(client):
    posts, _, _ = _import_posts(include_all=True)
    for post in posts:
        response = client.get(post.get_absolute_url())
        content = response.content.decode("utf-8")
        assert 'name="robots" content="index, follow"' in content


def test_import_command_idempotent():
    posts, titles, _ = _import_posts(include_all=True)
    first_count = posts.count()
    call_command("import_blog_content", "--all")
    posts_after = BlogPost.objects.filter(title__in=titles)
    assert posts_after.count() == first_count
    assert len(set(posts_after.values_list("slug", flat=True))) == len(titles)


def test_imported_post_content_quality():
    _, titles, data = _import_posts(include_all=True)
    assert all(title and title.strip() for title in titles)
    for item in data:
        assert item["title"].lower().endswith(".docx") is False
        assert 20 < len(item["title"]) <= 80
        assert 140 <= len(item["description"]) <= 240
        assert "<h2>" in item["content_html"]
        assert item["content_html"].count("<p>") >= 2


def test_import_default_filters_legacy_files(capsys):
    call_command("import_blog_content", "--dry-run")
    out = capsys.readouterr().out
    assert "Skipped (legacy)" in out
    assert "carfst_" in out


def test_import_with_all_includes_legacy(capsys):
    call_command("import_blog_content", "--dry-run", "--all")
    out = capsys.readouterr().out
    assert "Skipped (legacy)" not in out


def test_explicit_slug_respected(capsys):
    content_dir = Path(__file__).resolve().parents[1] / "_content"
    test_path = content_dir / "carfst_test_slug.docx"
    try:
        import docx

        doc = docx.Document()
        doc.add_paragraph("Slug: custom-slug-001")
        doc.add_paragraph("Title: Тестовая статья по пневмосистеме")
        doc.add_paragraph("Description: Короткое описание для теста импорта.")
        doc.add_heading("Вступление", level=2)
        doc.add_paragraph("Первый абзац текста.")
        doc.save(str(test_path))

        call_command("import_blog_content", "--dry-run")
        out = capsys.readouterr().out
        assert "custom-slug-001" in out
    finally:
        if test_path.exists():
            test_path.unlink()
